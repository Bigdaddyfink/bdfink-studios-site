"""
Generate BDFink Studios Resume and Cover Letter as Word documents.
Brand colors and typography aligned with brand-guidelines.html.
"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml
import os

# ─── Brand Colors ───
EMBER = RGBColor(0xE8, 0x55, 0x3D)
MIDNIGHT = RGBColor(0x0D, 0x0F, 0x12)
CHARCOAL = RGBColor(0x1A, 0x1D, 0x23)
SLATE = RGBColor(0x2A, 0x2E, 0x36)
SMOKE = RGBColor(0x8B, 0x8F, 0x96)
ASH = RGBColor(0xB4, 0xB8, 0xBF)
WHITE = RGBColor(0xF5, 0xF5, 0xF3)

FONT_BODY = "Calibri"
FONT_HEADING = "Georgia"

OUTPUT_DIR = os.path.dirname(os.path.abspath(__file__))


def set_cell_shading(cell, color_hex):
    """Set background color of a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_thin_line(doc):
    """Add a thin horizontal rule."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="4" w:space="1" w:color="E8553D"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def add_ember_line(doc):
    """Add a thick ember-colored accent line."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._element.get_or_add_pPr()
    pBdr = parse_xml(
        f'<w:pBdr {nsdecls("w")}>'
        f'<w:bottom w:val="single" w:sz="18" w:space="1" w:color="E8553D"/>'
        f'</w:pBdr>'
    )
    pPr.append(pBdr)


def style_run(run, font_name=FONT_BODY, size=11, bold=False, italic=False, color=MIDNIGHT):
    """Apply consistent styling to a run."""
    run.font.name = font_name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    run.font.color.rgb = color


def add_heading_styled(doc, text, level=1):
    """Add a styled heading with brand fonts."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16 if level == 1 else 10)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(text)
    if level == 1:
        style_run(run, FONT_HEADING, 14, bold=True, color=MIDNIGHT)
    elif level == 2:
        style_run(run, FONT_HEADING, 12, bold=True, color=SLATE)
    else:
        style_run(run, FONT_BODY, 11, bold=True, color=MIDNIGHT)
    return p


def add_body(doc, text, space_after=4, bold=False, italic=False, color=MIDNIGHT):
    """Add body text paragraph."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(14)
    run = p.add_run(text)
    style_run(run, FONT_BODY, 10.5, bold=bold, italic=italic, color=color)
    return p


def add_bullet(doc, text, color=MIDNIGHT):
    """Add a bullet point."""
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.line_spacing = Pt(13)
    # Clear default run and add styled one
    p.clear()
    run = p.add_run(text)
    style_run(run, FONT_BODY, 10, color=color)
    return p


def set_margins(doc, top=0.6, bottom=0.5, left=0.7, right=0.7):
    """Set page margins in inches."""
    for section in doc.sections:
        section.top_margin = Inches(top)
        section.bottom_margin = Inches(bottom)
        section.left_margin = Inches(left)
        section.right_margin = Inches(right)


# ═══════════════════════════════════════════════════════════════
#  RESUME
# ═══════════════════════════════════════════════════════════════

def generate_resume():
    doc = Document()
    set_margins(doc)

    # ─── HEADER ───
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TIMMY HARRIS")
    style_run(run, FONT_HEADING, 22, bold=True, color=MIDNIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Video Editor  ·  Motion Graphics Designer  ·  Illustrator")
    style_run(run, FONT_BODY, 11, color=EMBER)

    add_ember_line(doc)

    # Contact info in a compact line
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    contacts = [
        "Anderson, SC 29621",
        "770-713-1374",
        "bdfinkstudios@gmail.com",
    ]
    run = p.add_run("  ·  ".join(contacts))
    style_run(run, FONT_BODY, 9, color=SMOKE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    links = [
        "linkedin.com/in/bdfinkstudios",
        "bdfinkstudios.com",
        "art.bdfinkstudios.com",
    ]
    run = p.add_run("  ·  ".join(links))
    style_run(run, FONT_BODY, 9, color=SMOKE)

    add_thin_line(doc)

    # ─── SUMMARY ───
    add_body(
        doc,
        "SCAD-trained creative professional with 26+ years of experience in video editing, "
        "motion graphics, animation, illustration, and graphic design. Proven track record "
        "delivering results for Fortune 500 companies including IBM and Techtronic Industries. "
        "Currently producing video content for multiple clients while expanding into generative AI "
        "workflows for image and video creation.",
        space_after=6
    )

    add_thin_line(doc)

    # ─── CORE SKILLS ───
    h = add_heading_styled(doc, "CORE SKILLS", level=1)

    # Skills in grouped format
    skills_groups = [
        ("Video & Motion", "Video Editing, Motion Graphics, Animation, Color Correction, Sound Design, Audio Editing"),
        ("Design & Illustration", "Graphic Design, Illustration, Typography, Print & Collateral, Gig Poster Design"),
        ("Generative AI", "AI Image & Video Creation, Prompt Engineering, Midjourney, Kling, Luma, Stable Diffusion, Suno"),
        ("Management", "Project Management, Timeline & Resource Management, Stakeholder Collaboration"),
    ]
    for group_name, skills in skills_groups:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(13)
        run = p.add_run(f"{group_name}: ")
        style_run(run, FONT_BODY, 10, bold=True, color=MIDNIGHT)
        run = p.add_run(skills)
        style_run(run, FONT_BODY, 10, color=SLATE)

    add_thin_line(doc)

    # ─── PROFESSIONAL EXPERIENCE ───
    add_heading_styled(doc, "PROFESSIONAL EXPERIENCE", level=1)

    # --- BClip Productions ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BClip Productions")
    style_run(run, FONT_HEADING, 11.5, bold=True, color=MIDNIGHT)
    run = p.add_run("  —  Asheville, NC (Remote)")
    style_run(run, FONT_BODY, 10, color=SMOKE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Video Editor & Motion Graphics Designer")
    style_run(run, FONT_BODY, 10, italic=True, color=SLATE)
    run = p.add_run("  |  Feb 2025 – Present")
    style_run(run, FONT_BODY, 9, color=SMOKE)

    bullets_bclip = [
        "Edit and produce marketing video content for social media and broadcast distribution, managing projects from raw footage to final delivery.",
        "Create on-screen motion graphics, animated titles, and visual effects to elevate storytelling impact.",
        "Perform color correction, voice-over editing, and music editing to deliver polished final products.",
        "Collaborate with Art Director on script development, visual concepts, and pacing to meet client timelines.",
        "Research and repurpose existing video assets to maximize content value across campaigns.",
    ]
    for b in bullets_bclip:
        add_bullet(doc, b)

    # --- Well AD ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Well AD")
    style_run(run, FONT_HEADING, 11.5, bold=True, color=MIDNIGHT)
    run = p.add_run("  —  Los Angeles, CA (Remote)")
    style_run(run, FONT_BODY, 10, color=SMOKE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Video Editor & Motion Graphics Designer")
    style_run(run, FONT_BODY, 10, italic=True, color=SLATE)
    run = p.add_run("  |  Dec 2024 – Present")
    style_run(run, FONT_BODY, 9, color=SMOKE)

    bullets_wellad = [
        "Produce short-form Instagram Reels for restaurant marketing campaigns across the Los Angeles market.",
        "Write scripts, design hooks, and create attention-grabbing intros optimized for social media engagement.",
        "Handle full post-production pipeline: editing, color correction, motion graphics, closed captioning, and audio enhancement.",
        "Deliver rapid-turnaround content while maintaining consistent brand quality across multiple client accounts.",
    ]
    for b in bullets_wellad:
        add_bullet(doc, b)

    # --- Starboard ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Starboard")
    style_run(run, FONT_HEADING, 11.5, bold=True, color=MIDNIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Video Editor & Motion Graphics Designer")
    style_run(run, FONT_BODY, 10, italic=True, color=SLATE)
    run = p.add_run("  |  Jul 2023 – Oct 2023")
    style_run(run, FONT_BODY, 9, color=SMOKE)

    bullets_starboard = [
        "Developed marketing video content for social media, YouTube, and online advertising campaigns.",
        "Pioneered use of generative AI tools to create visual assets and concept art for client projects.",
        "Conducted audio editing and voice-over enhancements for final deliverables.",
        "Managed concurrent projects in a fast-paced agency environment, meeting all client deadlines.",
    ]
    for b in bullets_starboard:
        add_bullet(doc, b)

    # --- TTI ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Techtronic Industries (TTI) — Torque Creative")
    style_run(run, FONT_HEADING, 11.5, bold=True, color=MIDNIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Video Editor, Motion Graphics Specialist & Project Manager")
    style_run(run, FONT_BODY, 10, italic=True, color=SLATE)
    run = p.add_run("  |  Dec 2019 – Nov 2022")
    style_run(run, FONT_BODY, 9, color=SMOKE)

    bullets_tti = [
        "Led development, production, and post-production of marketing video content for Fortune 500 product launches (Milwaukee, Ryobi, RIDGID brands).",
        "Produced on-screen graphics and edited videos from raw footage to final deliverables across multiple formats.",
        "Managed project timelines and coordinated with internal teams and external vendors to ensure on-time delivery.",
        "Collaborated with editors to maintain brand standards and quality across all video assets.",
        "Created animated PowerPoint presentations with embedded video for sales teams and executive leadership.",
        "Organized and archived raw video footage and project files for future use.",
    ]
    for b in bullets_tti:
        add_bullet(doc, b)

    # --- IBM ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("IBM — Americas Learning Development")
    style_run(run, FONT_HEADING, 11.5, bold=True, color=MIDNIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Media Specialist / Senior Animator")
    style_run(run, FONT_BODY, 10, italic=True, color=SLATE)
    run = p.add_run("  |  2003 – 2019")
    style_run(run, FONT_BODY, 9, color=SMOKE)

    bullets_ibm = [
        "Led multimedia design and integration for enterprise education and training programs over 16 years.",
        "Specialized in graphic animation, simulations, audio/video production, and interactive Flash/HTML content.",
        "Designed product identities, logos, marketing materials, and custom packaging for retail displays.",
        "Created animated PowerPoint presentations with embedded video for internal and external audiences.",
        "Mentored junior designers and established multimedia production workflows and standards.",
    ]
    for b in bullets_ibm:
        add_bullet(doc, b)

    # --- BDFink Studios ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BDFink Studios")
    style_run(run, FONT_HEADING, 11.5, bold=True, color=MIDNIGHT)
    run = p.add_run("  —  Freelance")
    style_run(run, FONT_BODY, 10, color=SMOKE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run("Illustration, Motion Graphics, Video Editing & Graphic Design")
    style_run(run, FONT_BODY, 10, italic=True, color=SLATE)
    run = p.add_run("  |  2000 – Present")
    style_run(run, FONT_BODY, 9, color=SMOKE)

    bullets_bd = [
        "Operate a freelance creative studio specializing in video editing, motion graphics, animation, illustration, and graphic design.",
        "Create hand-drawn illustrations, gig posters, and limited-edition print media for bands, events, and private commissions.",
        "Produce digital video content including commercials, social media shorts, training videos, and trade show presentations.",
        "Deliver custom motion graphics, color grading, audio editing, sound design, and animated logos.",
        "Leverage generative AI tools (Midjourney, Kling, Luma, Stable Diffusion, ComfyUI, Suno) for image and video creation, pre-visualization, and content development.",
    ]
    for b in bullets_bd:
        add_bullet(doc, b)

    add_thin_line(doc)

    # ─── TECHNICAL PROFICIENCIES ───
    add_heading_styled(doc, "TECHNICAL PROFICIENCIES", level=1)

    tech_groups = [
        ("Video & Motion", "After Effects, Premiere Pro, DaVinci Resolve, Media Encoder, Red Giant, Animation Composer, CapCut, Camtasia"),
        ("Design & Illustration", "Photoshop, Illustrator, InDesign, Adobe Animate, Canva, Adobe Express, Figma"),
        ("Generative AI", "Midjourney, Stable Diffusion, Leonardo AI, Adobe Firefly, Runway, Kling AI, Luma Labs, Suno AI, ComfyUI, Pinokio, n8n"),
        ("Collaboration", "Workfront, Monday, Trello, Frame.io, Slack, Dropbox, Google Drive, Audacity, PowerPoint"),
    ]
    for group_name, tools in tech_groups:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.line_spacing = Pt(13)
        run = p.add_run(f"{group_name}: ")
        style_run(run, FONT_BODY, 10, bold=True, color=MIDNIGHT)
        run = p.add_run(tools)
        style_run(run, FONT_BODY, 10, color=SLATE)

    add_thin_line(doc)

    # ─── EDUCATION ───
    add_heading_styled(doc, "EDUCATION", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Bachelor of Fine Arts")
    style_run(run, FONT_HEADING, 11, bold=True, color=MIDNIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("The Savannah College of Art and Design (SCAD)  —  Savannah, Georgia")
    style_run(run, FONT_BODY, 10, color=SLATE)

    add_thin_line(doc)

    # ─── INDUSTRIES ───
    add_heading_styled(doc, "INDUSTRIES", level=1)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run(
        "Video Production  ·  Post-Production  ·  Marketing  ·  Education & Training  ·  "
        "Digital & Print Media  ·  Packaging  ·  Entertainment  ·  Generative AI"
    )
    style_run(run, FONT_BODY, 10, color=SLATE)

    # Save
    path = os.path.join(OUTPUT_DIR, "Timmy_Harris_Resume.docx")
    doc.save(path)
    print(f"Resume saved: {path}")
    return path


# ═══════════════════════════════════════════════════════════════
#  COVER LETTER
# ═══════════════════════════════════════════════════════════════

def generate_cover_letter():
    doc = Document()
    set_margins(doc, top=1.0, bottom=1.0, left=1.0, right=1.0)

    # ─── HEADER ───
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("TIMMY HARRIS")
    style_run(run, FONT_HEADING, 20, bold=True, color=MIDNIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Video Editor  ·  Motion Graphics Designer  ·  Illustrator")
    style_run(run, FONT_BODY, 10, color=EMBER)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("Anderson, SC 29621  ·  770-713-1374  ·  bdfinkstudios@gmail.com")
    style_run(run, FONT_BODY, 9, color=SMOKE)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run("bdfinkstudios.com  ·  linkedin.com/in/bdfinkstudios")
    style_run(run, FONT_BODY, 9, color=SMOKE)

    add_ember_line(doc)

    # ─── DATE & ADDRESS ───
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run("[DATE]")
    style_run(run, FONT_BODY, 10.5, color=MIDNIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("[HIRING MANAGER NAME]")
    style_run(run, FONT_BODY, 10.5, color=MIDNIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("[COMPANY NAME]")
    style_run(run, FONT_BODY, 10.5, color=MIDNIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(16)
    run = p.add_run("[COMPANY ADDRESS]")
    style_run(run, FONT_BODY, 10.5, color=MIDNIGHT)

    # ─── GREETING ───
    add_body(doc, "Dear [HIRING MANAGER NAME],", space_after=12)

    # ─── OPENING ───
    add_body(
        doc,
        "I'm writing to apply for the [ROLE TITLE] position at [COMPANY NAME]. "
        "With 26+ years of experience in video editing, motion graphics, animation, and graphic design — "
        "including 16 years at IBM and post-production leadership at Techtronic Industries — I bring both "
        "the creative range and the production discipline your team needs. "
        "[SPECIFIC THING YOU ADMIRE ABOUT THE COMPANY OR ROLE — e.g., "
        "\"Your recent campaign for X caught my attention because...\"]",
        space_after=10
    )

    # ─── WHAT I BRING ───
    add_body(
        doc,
        "I currently work as a remote contractor for two companies simultaneously: bClip Productions "
        "in Asheville, NC, where I edit and produce storytelling-driven marketing videos, and Well AD in "
        "Los Angeles, where I create short-form Instagram Reels for their restaurant marketing clients. "
        "I handle the full post-production pipeline on every project — editing, motion graphics, color "
        "correction, audio, and delivery — and I consistently meet tight deadlines across multiple accounts.",
        space_after=10
    )

    # ─── WHAT SETS ME APART ───
    add_body(
        doc,
        "What sets me apart is the breadth of my creative background. My BFA from the Savannah College "
        "of Art and Design grounded me in illustration and design fundamentals that still inform every "
        "frame I edit. I'm as comfortable building motion graphics packages in After Effects as I am "
        "hand-drawing a gig poster. I've also been actively integrating generative AI tools — Midjourney, "
        "Kling, Luma, Stable Diffusion, and others — into my creative workflow, using them for "
        "pre-visualization, asset creation, and content development.",
        space_after=10
    )

    # ─── FIT ───
    add_body(
        doc,
        "I thrive in collaborative, deadline-driven environments. At TTI, I managed project timelines "
        "across internal teams and external vendors for Fortune 500 product launches. At IBM, I led "
        "multimedia production for enterprise training programs for over a decade. I'm experienced with "
        "remote workflows and the collaboration tools that make them work — Slack, Frame.io, Workfront, "
        "Monday, and Dropbox.",
        space_after=10
    )

    # ─── CLOSING ───
    add_body(
        doc,
        "I'd welcome the opportunity to discuss how my experience and skills align with what you're "
        "looking for. You can see examples of my recent work at bdfinkstudios.com — including current "
        "video editing projects (bdfinkstudios.com/video-editing-work-examples-2025) and Instagram "
        "Reels (bdfinkstudios.com/instagram-reels-examples).",
        space_after=10
    )

    add_body(doc, "Thank you for your time and consideration.", space_after=20)

    add_body(doc, "Best regards,", space_after=4)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("Timmy Harris")
    style_run(run, FONT_HEADING, 12, bold=True, color=MIDNIGHT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run("BDFink Studios")
    style_run(run, FONT_BODY, 10, color=SLATE)

    # Save
    path = os.path.join(OUTPUT_DIR, "Timmy_Harris_Cover_Letter.docx")
    doc.save(path)
    print(f"Cover letter saved: {path}")
    return path


if __name__ == "__main__":
    print("Generating BDFink Studios documents...\n")
    generate_resume()
    generate_cover_letter()
    print("\nDone! Open the .docx files in Word to review.")
